"""
Conversor Markdown → .docx genérico para los reportes de Contexto AI.
Maneja: títulos, tablas, código en bloque, citas, viñetas, listas numeradas,
**negritas** e `inline code` reales.

Uso:
    python docs/_md2docx_generic.py docs/REPORTE_X.md  [nombre_salida.docx]
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

TEAL = RGBColor(0x1A, 0x7A, 0x76)
CODE_BLUE = RGBColor(0x1F, 0x6F, 0xEB)


def add_runs_with_markup(p, text):
    """Agrega 'text' a un parrafo interpretando **negrita** e `inline code`."""
    # Tokeniza por **...** y `...`
    parts = re.split(r"(\*\*.+?\*\*|`.+?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9.5)
            r.font.color.rgb = CODE_BLUE
        else:
            p.add_run(part)


def convert(src_path: Path, out_paths: list[Path]):
    src = src_path.read_text(encoding="utf-8")
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    lines = src.split("\n")
    in_code = False
    code_buf: list[str] = []
    tbl_rows: list[str] = []

    def flush_table():
        nonlocal tbl_rows
        rows = [r for r in tbl_rows if not re.match(r"^\s*\|[\s:|-]+\|\s*$", r)]
        cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
        if cells:
            t = doc.add_table(rows=len(cells), cols=len(cells[0]))
            t.style = "Light Grid Accent 1"
            for i, row in enumerate(cells):
                for j, val in enumerate(row):
                    if j < len(t.rows[i].cells):
                        cell = t.rows[i].cells[j]
                        cell.text = ""
                        add_runs_with_markup(cell.paragraphs[0], val)
                        if i == 0:
                            for rr in cell.paragraphs[0].runs:
                                rr.bold = True
        tbl_rows = []

    def add_code(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = "Consolas"; r.font.size = Pt(9); r.font.color.rgb = CODE_BLUE

    for ln in lines:
        if ln.strip().startswith("```"):
            if in_code:
                add_code("\n".join(code_buf)); code_buf = []; in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_buf.append(ln); continue

        if ln.strip().startswith("|"):
            tbl_rows.append(ln); continue
        elif tbl_rows:
            flush_table()

        s = ln.rstrip()
        stripped = s.strip()

        if s.startswith("# "):
            doc.add_heading(s[2:], level=0)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=1)
        elif s.startswith("### "):
            doc.add_heading(s[4:], level=2)
        elif stripped.startswith(">"):
            p = doc.add_paragraph(); p.style = "Intense Quote"
            add_runs_with_markup(p, stripped.lstrip("> ").strip())
        elif re.match(r"^\s*[-*]\s+", s):
            indent = len(s) - len(s.lstrip())
            style = "List Bullet 2" if indent >= 2 else "List Bullet"
            p = doc.add_paragraph(style=style)
            add_runs_with_markup(p, re.sub(r"^\s*[-*]\s+", "", s))
        elif re.match(r"^\s*\d+\.\s+", s):
            p = doc.add_paragraph(style="List Number")
            add_runs_with_markup(p, re.sub(r"^\s*\d+\.\s+", "", s))
        elif stripped.startswith("---") or stripped == "":
            continue
        else:
            p = doc.add_paragraph()
            add_runs_with_markup(p, s)

    if tbl_rows:
        flush_table()

    for out in out_paths:
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))
        print("OK ->", out)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Uso: python docs/_md2docx_generic.py <archivo.md> [salida.docx]"); sys.exit(1)
    src = Path(sys.argv[1]).resolve()
    name = sys.argv[2] if len(sys.argv) > 2 else src.stem + ".docx"
    outs = [
        Path(r"C:\Users\DETPC\Desktop\Contexxto") / name,
        src.parent / name,  # junto al .md (docs/)
    ]
    convert(src, outs)
