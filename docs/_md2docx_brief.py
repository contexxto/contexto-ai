"""Convierte el diseño Brief Intake Multimodal a .docx para Gemini."""
import re
from docx import Document
from docx.shared import Pt, RGBColor

src = open("docs/DISENO_Brief_Intake_Multimodal.md", encoding="utf-8").read()
doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


def add_code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1F, 0x6F, 0xEB)


def strip_md(s):
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    return s.replace("`", "")


lines = src.split("\n")
in_code = False
code_buf = []
tbl_rows = []


def flush_table():
    global tbl_rows
    rows = [r for r in tbl_rows if not re.match(r"^\s*\|[\s:|-]+\|\s*$", r)]
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    if cells:
        t = doc.add_table(rows=len(cells), cols=len(cells[0]))
        t.style = "Light Grid Accent 1"
        for i, row in enumerate(cells):
            for j, val in enumerate(row):
                if j < len(t.rows[i].cells):
                    t.rows[i].cells[j].text = strip_md(val)
    tbl_rows = []


for ln in lines:
    if ln.strip().startswith("```"):
        if in_code:
            add_code("\n".join(code_buf)); code_buf = []; in_code = False
        else:
            in_code = True
        continue
    if in_code:
        code_buf.append(ln); continue
    if ln.strip().startswith("|"):
        tbl_rows.append(ln); continue
    elif tbl_rows:
        flush_table()
    s = ln.rstrip()
    if s.startswith("# "):
        doc.add_heading(strip_md(s[2:]), level=0)
    elif s.startswith("## "):
        doc.add_heading(strip_md(s[3:]), level=1)
    elif s.startswith("### "):
        doc.add_heading(strip_md(s[4:]), level=2)
    elif s.startswith(">"):
        p = doc.add_paragraph(strip_md(s.lstrip("> ").strip())); p.style = "Intense Quote"
    elif s.startswith("---") or s.strip() == "":
        continue
    else:
        doc.add_paragraph(strip_md(s))

if tbl_rows:
    flush_table()

for out in [r"C:\Users\DETPC\Desktop\Contexxto\DISENO_Brief_Intake_Multimodal.docx",
            r"C:\Users\DETPC\Desktop\Contexto-AI\docs\DISENO_Brief_Intake_Multimodal.docx"]:
    doc.save(out)
    print("OK ->", out)
